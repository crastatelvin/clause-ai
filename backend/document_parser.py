# ============================================
# Project: CLAUSE — AI Contract Risk Analyzer
# Author: Telvin Crasta
# GitHub: github.com/crastatelvin
# License: CC BY-NC 4.0
# Original design and architecture by Telvin Crasta
# ============================================

from __future__ import annotations

import io
import logging

import fitz  # PyMuPDF

log = logging.getLogger(__name__)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract digital-text layer from a PDF. Empty string if scan-only."""
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    try:
        return "\n".join(page.get_text() for page in doc).strip()
    finally:
        doc.close()


def extract_text_from_txt(file_bytes: bytes) -> str:
    return file_bytes.decode("utf-8", errors="ignore").strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a Word .docx using python-docx. Imported lazily so
    the rest of the pipeline still works if the dependency is missing."""
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            "DOCX support requires python-docx. Run `pip install python-docx`."
        ) from exc

    doc = Document(io.BytesIO(file_bytes))
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        if paragraph.text:
            parts.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    parts.append(text)
    return "\n".join(parts).strip()


def _detect_file_type(filename: str) -> str:
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return "pdf"
    if lower.endswith(".docx"):
        return "docx"
    if lower.endswith((".txt", ".md")):
        return "text"
    return "text"  # best-effort fallback


def parse_document(file_bytes: bytes, filename: str) -> dict:
    file_type = _detect_file_type(filename)

    if file_type == "pdf":
        text = extract_text_from_pdf(file_bytes)
    elif file_type == "docx":
        text = extract_text_from_docx(file_bytes)
    else:
        text = extract_text_from_txt(file_bytes)

    word_count = len(text.split())
    return {
        "text": text,
        "file_type": file_type,
        "word_count": word_count,
        "char_count": len(text),
        "filename": filename,
    }
